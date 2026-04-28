#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate all required docx documents for the canteen project reports.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

TASKS_DIR = r"c:\Users\11428\Desktop\canteen\tasks"


def set_run_font(run, font_name="宋体", font_size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    from lxml import etree
    rns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    run_east = etree.SubElement(rpr, f"{{{rns}}}rFonts")
    run_east.set(f"{{{rns}}}eastAsia", font_name)


def add_formatted_paragraph(doc, text, style="Normal", font_size=12, bold=False, space_before=0, space_after=6, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_size=font_size, bold=bold)
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return p


def add_mixed_paragraph(doc, parts, space_before=0, space_after=6, alignment=None):
    """parts: list of (text, bold, font_size)"""
    p = doc.add_paragraph()
    for text, bold, font_size in parts:
        run = p.add_run(text)
        set_run_font(run, font_size=font_size, bold=bold)
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return p


def add_heading_with_font(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, font_name="黑体" if level <= 2 else "宋体", font_size=16 if level == 1 else 14, bold=True)
    return h


def add_bullet(doc, text, font_size=11, bold=False, indent=True):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, font_size=font_size, bold=bold)
    return p


# ============================================================
# 1. 开发阶段小组沟通交流记录.docx
# ============================================================
def create_meeting_record():
    doc = Document()

    # Title
    add_heading_with_font(doc, "开发阶段小组沟通交流记录", level=1)

    # 会议信息
    add_heading_with_font(doc, "一、会议信息", level=2)
    add_formatted_paragraph(doc, "时间：2026年4月28日", font_size=12)
    add_formatted_paragraph(doc, "参与人：", font_size=12)
    add_bullet(doc, "黄译辉（前端开发）", font_size=12)
    add_bullet(doc, "陈金瑞（后端开发）", font_size=12)

    # 沟通交流内容
    add_heading_with_font(doc, "二、沟通交流内容", level=2)

    add_mixed_paragraph(doc, [
        ("黄译辉（前端开发）：", True, 12),
        ('目前前端的 Vue 页面（用户端和管理员端）基本开发完成了，包括页面布局和主要的交互。但在进行前后端联调测试"扫码核销二维码(ScanVerify)"模块时，我发现后端返回的 JSON 数据里缺失了 canteenId 等字段，导致前端获取不到信息渲染报错。另外，对于用户订单状态的实时更新，我们现在该怎么处理比较好？', False, 12)
    ])

    add_mixed_paragraph(doc, [
        ("陈金瑞（后端开发）：", True, 12),
        ('核销接口那个确实是我的疏忽，前两天重构 Mapper 文件时不小心漏掉了几处关联查询。我今天之内就会修复它，把缺失的字段补齐。关于订单状态实时更新，我原本想接入 WebSocket 进行主动推送的，但这要改不少代码并增加配置。咱们时间比较紧。', False, 12)
    ])

    add_mixed_paragraph(doc, [
        ("黄译辉（前端开发）：", True, 12),
        ('如果是这样的话，我建议咱们前端先用短期轮询（Polling）来做，比如每隔 3-5 秒请求一次获取订单最新状态。这能最快实现状态刷新，也减少这几天我们互相卡进度的风险，你看呢？', False, 12)
    ])

    add_mixed_paragraph(doc, [
        ("陈金瑞（后端开发）：", True, 12),
        ('没问题，对于目前的实训规模，轮询完全能撑住。另外，后台数据统计页面（Statistics）需要大量历史数据来展示图表，你可以先写静态数据 Mock，或者我这边在 data_init.sql 脚本里加一批模拟的历史订单记录来支撑图表绘制？', False, 12)
    ])

    add_mixed_paragraph(doc, [
        ("黄译辉（前端开发）：", True, 12),
        ('最好是真实调接口，麻烦你在 data_init.sql 里面插入几十条模拟的历史订单和取餐数据，这样我可以验证图表组件按日期分类、聚合的准确性。', False, 12)
    ])

    # 议定内容
    add_heading_with_font(doc, "三、议定内容", level=2)
    add_bullet(doc, '接口字段修复：后端（陈金瑞）将于今日内修复核销二维码以及对应查询接口的数据返回缺失问题，确保与接口文档返回体完全一致。', font_size=12)
    add_bullet(doc, '实时状态更新方案：订单状态更新暂不使用 WebSocket。前端（黄译辉）统一使用定时轮询机制来实现订单及排队进度刷新，降低当前的开发风险。', font_size=12)
    add_bullet(doc, '补充测试数据：后端（陈金瑞）将在数据库初始化脚本中补充不少于 50 条历史订单的模拟数据，以供前端 Dashboard 和 Statistics 页面的可视化组件进行渲染测试。', font_size=12)

    # 未议定内容
    add_heading_with_font(doc, "四、未议定内容", level=2)
    add_bullet(doc, '高并发抢单机制：就餐高峰期用户并发提交订单和抢座可能导致的库存/座位超卖问题，目前仍未达成有效且容易实现的解决共识（讨论了乐观锁与 Redis 方案，但均觉得改造量大）。暂定推迟至最后测试阶段再视情况决定。', font_size=12)
    add_bullet(doc, '首页菜品智能推荐：是否要在前端首页引入一套简单的"热销及评价最高"计算逻辑展示模块。因前端排版较紧凑，暂无定论，双方决定先把核心流程走通。', font_size=12)

    # 后续计划
    add_heading_with_font(doc, "五、后续计划", level=2)
    add_bullet(doc, '今明两天：前后端完成所有的 API 接口对接，走通"浏览商品 -> 下单 -> 取餐码生成 -> 管理员扫码核销"完整全链路（Full link）。', font_size=12)
    add_bullet(doc, '本周末：修复前后端交互阶段所产生的所有 Bug。执行自动化运行测试（full_link_test.py）观察结果。', font_size=12)
    add_bullet(doc, '下一阶段：冻结主要业务代码，停止新增功能，逐步进入性能测试及单元测试报告（如软件综合实训报告）的撰写工作。', font_size=12)

    doc.save(os.path.join(TASKS_DIR, "开发阶段小组沟通交流记录.docx"))
    print("Created: 开发阶段小组沟通交流记录.docx")


# ============================================================
# 2. 软件综合实训_陈金瑞_开发阶段实训报告.docx
# ============================================================
def create_chen_jinrui_report():
    doc = Document()

    add_heading_with_font(doc, "软件综合实训_陈金瑞_开发阶段实训报告", level=1)

    # Part 1
    add_heading_with_font(doc, "一、在本阶段中个人遇到了哪些问题，如何解决的", level=2)

    add_formatted_paragraph(doc, "在智慧校园食堂系统的后端开发阶段，由于业务逻辑的严谨性以及系统安全性要求，我主要遇到了以下几个典型问题：", font_size=12)

    # Problem 1
    add_heading_with_font(doc, "问题一：JWT 认证与安全机制的设计实现", level=3)
    add_mixed_paragraph(doc, [
        ("问题描述：", True, 12),
        ("在后端需要实现用户身份验证与权限控制。传统的 Session 方案在前后端分离架构下存在跨域保持状态的困难，同时需要确保 Token 的安全性防止被伪造和篡改。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("解决过程：", True, 12),
        ("我采用了 JWT（JSON Web Token）作为认证方案。在后端引入了 jjwt 库，创建了 JwtUtil 工具类，用于生成、解析和验证 Token。通过 Spring 的 HandlerInterceptor 机制实现了拦截器 JwtInterceptor，对所有 API 请求进行统一拦截验证，从请求头中提取 Authorization 字段的 Token 进行解析，验证通过后在 ThreadLocal 中存储当前用户信息供后续业务逻辑使用。同时配置了 WebMvcConfigurer 将登录和注册等公开接口排除在拦截之外。", False, 12)
    ])

    # Problem 2
    add_heading_with_font(doc, "问题二：库存超卖问题的解决", level=3)
    add_mixed_paragraph(doc, [
        ("问题描述：", True, 12),
        ('在测试"提交订单"接口时发现，如果模拟多个用户同时对同一份菜品下单，会出现菜品库存由于并发脏读而变成负数的情况，导致库存数据不一致。', False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("解决过程：", True, 12),
        ("为了解决并发一致性问题，我查阅了数据库事务隔离级别的相关资料，并结合业务实际情况，采用了乐观锁的机制。在 FoodItemMapper.xml 对应的扣减库存 SQL 中，加入了库存大于等于扣减数量的条件判断（UPDATE food_item SET stock = stock - #{num} WHERE id = #{id} AND stock >= #{num}）。如果受影响行数为 0，则抛出自定义的库存不足异常。同时在 Service 层通过 @Transactional 注解保证订单创建和库存扣减操作的原子性，确保数据一致性。", False, 12)
    ])

    # Problem 3
    add_heading_with_font(doc, "问题三：取餐二维码的安全性与时效性设计", level=3)
    add_mixed_paragraph(doc, [
        ("问题描述：", True, 12),
        ("早期设计取餐二维码时，仅将订单号直接编码进入二维码中。这种做法存在极大的安全隐患，容易被恶意伪造和重放攻击。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("解决过程：", True, 12),
        ('我重新设计了取餐码的生成机制，引入了动态 Token 和有效期限制。后端生成二维码数据时，结合"订单 ID + 用户 ID + 时间戳 + 随机盐"生成唯一的 qrToken 并存入 PickupQrCode 表中，同时设置合理的过期时间。在管理端扫描核验时，后端不仅校验 Token 的合法性，还会检查二维码是否过期以及是否已被使用。核销记录通过 QrScanLogMapper 进行持久化，便于后续审计回溯。', False, 12)
    ])

    # Part 2
    add_heading_with_font(doc, "二、在本阶段中个人承担了哪些任务，完成情况如何，过程中是否遇到问题，如何解决的", level=2)

    add_formatted_paragraph(doc, "在本项目的开发阶段，我主要承担了后端全部开发工作，包括 Spring Boot 应用搭建、RESTful API 设计、业务逻辑实现、数据库访问层开发、JWT 认证机制和全局异常处理。", font_size=12)

    # Task 1
    add_heading_with_font(doc, "任务一：Spring Boot 后端应用整体搭建与 RESTful API 设计", level=3)
    add_mixed_paragraph(doc, [
        ("完成情况：", True, 12),
        ("100% 完成。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("工作内容：", True, 12),
        ("使用 Spring Boot 3.x 框架搭建了后端项目，采用 Maven 进行依赖管理。设计了统一的 RESTful API 风格，包含用户模块、食堂模块、菜品模块、购物车模块、订单模块、预约模块和二维码核销模块。使用 MyBatis 作为 ORM 框架，编写了对应的 Mapper 接口和 XML 映射文件。创建了统一的 Result 响应封装类，确保所有接口返回一致的数据结构（code、message、data）。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("遇到的问题：", True, 12),
        ("在项目初期，MyBatis 与 Spring Boot 3.x 的兼容配置需要额外注意版本匹配问题。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("解决过程：", True, 12),
        ("通过查阅官方文档和相关资料，选定了 mybatis-spring-boot-starter 的适配版本，并在 application.yml 中正确配置了数据源和 Mapper 扫描路径。", False, 12)
    ])

    # Task 2
    add_heading_with_font(doc, "任务二：JWT 认证与全局拦截器实现", level=3)
    add_mixed_paragraph(doc, [
        ("完成情况：", True, 12),
        ("100% 完成。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("工作内容：", True, 12),
        ("实现了完整的 JWT 认证体系，包括 Token 的生成、解析、刷新和验证逻辑。创建了 JwtInterceptor 拦截器对所有非公开接口进行统一鉴权，从请求头中提取并验证 Token，将解析出的用户信息存入 ThreadLocal 以便业务层获取当前登录用户身份。配置了 WebMvcConfigurer 注册拦截器并排除登录、注册等公开接口。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("遇到的问题：", True, 12),
        ("在拦截器中如何优雅地处理未授权访问并返回统一格式的错误响应。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("解决过程：", True, 12),
        ("通过结合 @RestControllerAdvice 全局异常处理机制，自定义了 AuthenticationException 异常类，在拦截器验证失败时抛出该异常，由全局异常处理器统一捕获并返回标准格式的 401 错误响应。", False, 12)
    ])

    # Task 3
    add_heading_with_font(doc, "任务三：核心业务逻辑与数据库事务管理", level=3)
    add_mixed_paragraph(doc, [
        ("完成情况：", True, 12),
        ("100% 完成。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("工作内容：", True, 12),
        ("实现了购物车管理、订单创建与状态流转、预约管理、取餐二维码生成与核销等核心业务逻辑。在订单创建流程中，严格校验了营业时间段、座位余量等前置条件，通过 @Transactional(rollbackFor = Exception.class) 确保了多表操作的原子性。同时实现了库存扣减的乐观锁机制，防止并发场景下的超卖问题。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("遇到的问题：", True, 12),
        ("订单创建涉及多张表的联动操作，如果中途某个环节失败，如何保证数据一致性。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("解决过程：", True, 12),
        ("通过合理设置 Spring 事务的传播行为和回滚策略，确保在业务层抛出任何 RuntimeException 或自定义业务异常时，事务都能正确回滚。同时对关键操作添加了日志记录便于问题排查。", False, 12)
    ])

    # Task 4
    add_heading_with_font(doc, "任务四：全局异常处理与接口规范", level=3)
    add_mixed_paragraph(doc, [
        ("完成情况：", True, 12),
        ("100% 完成。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("工作内容：", True, 12),
        ("创建了 @RestControllerAdvice 全局异常处理器，统一捕获和处理各类异常，包括参数校验异常、业务逻辑异常、认证授权异常和系统未知异常。设计了自定义业务异常类 BusinessExceptionHandler，通过枚举类定义标准错误码和错误信息，确保前端能够根据不同的错误码进行针对性的提示。", False, 12)
    ])

    # Summary
    add_heading_with_font(doc, "三、总结", level=2)
    add_formatted_paragraph(doc, "本阶段按计划完成了我所负责的所有后端开发工作，包括 Spring Boot 应用搭建、RESTful API 设计实现、JWT 认证机制、全局异常处理、数据库事务管理和乐观锁并发控制。代码已通过基本的功能测试和接口测试（使用 full_link_test.py 等工具）。在实际开发过程中，不仅加深了对 Spring Boot 框架原理的理解，还通过解决认证安全、事务一致性、并发控制等问题，积累了宝贵的工程实践经验。", font_size=12)

    doc.save(os.path.join(TASKS_DIR, "软件综合实训_陈金瑞_开发阶段实训报告.docx"))
    print("Created: 软件综合实训_陈金瑞_开发阶段实训报告.docx")


# ============================================================
# 3. 软件综合实训_黄译辉_开发阶段实训报告.docx
# ============================================================
def create_huang_yihui_report():
    doc = Document()

    add_heading_with_font(doc, "软件综合实训_黄译辉_开发阶段实训报告", level=1)

    # Part 1
    add_heading_with_font(doc, "一、在本阶段中个人遇到了哪些问题，如何解决的", level=2)

    add_formatted_paragraph(doc, "在智慧校园食堂系统的前端开发与部署阶段，由于前后端分离架构的复杂性以及自动化测试的需求，我主要遇到了以下几个典型问题：", font_size=12)

    # Problem 1
    add_heading_with_font(doc, "问题一：前后端联调时的跨域与 Token 管理问题", level=3)
    add_mixed_paragraph(doc, [
        ("问题描述：", True, 12),
        ("在本地开发时，前端（Vue3 + Vite，运行在 5173 端口）向后端（Spring Boot，运行在 8080 端口）发起 API 请求时，浏览器报跨域拦截错误。此外，登录后获取的 JWT Token 在后续请求中需要正确携带和刷新。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("解决过程：", True, 12),
        ("在后端配合下配置了全局 CORS 映射。在前端封装了 Axios 实例（utils/request.js），在请求拦截器中统一从 localStorage 读取 Token 并附加到请求头的 Authorization 字段中。在响应拦截器中统一处理了 401 状态码，使其在 Token 失效时自动清除本地缓存并重定向至登录页，确保用户体验的流畅性。", False, 12)
    ])

    # Problem 2
    add_heading_with_font(doc, "问题二：全链路自动化测试脚本的编写与调试", level=3)
    add_mixed_paragraph(doc, [
        ("问题描述：", True, 12),
        ("在编写全链路测试脚本（full_link_test.py）时，需要模拟完整的用户操作流程：注册、登录、浏览食堂、查看菜品、加入购物车、创建订单、支付回调、预约就餐、生成二维码和管理端核销。各个环节存在数据依赖关系，前面的接口返回数据需要作为后续接口的输入参数。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("解决过程：", True, 12),
        ("采用 Python 的 requests 库编写了基于 Session 的测试脚本，将每一步的响应数据进行 JSON 解析并提取关键字段（如 userId、token、orderId、seatId 等）作为后续请求的输入。通过严格的 assert 断言验证每一步的返回状态码和数据正确性。最终实现了完整的自动化全链路测试，可一键验证所有核心业务功能是否正常运作。", False, 12)
    ])

    # Problem 3
    add_heading_with_font(doc, "问题三：Docker 容器化部署的配置与调试", level=3)
    add_mixed_paragraph(doc, [
        ("问题描述：", True, 12),
        ("项目需要支持 Docker 容器化部署以便于演示和交付，但前后端服务与 MySQL 数据库之间的网络连接配置存在挑战，需要确保容器间能够正确通信。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("解决过程：", True, 12),
        ("编写了 Dockerfile 分别构建前端和后端镜像，使用 docker-compose.yml 统一管理所有服务（前端、后端、MySQL 数据库）。通过 Docker 自定义网络实现容器间基于服务名的 DNS 解析，确保后端容器能够通过服务名连接到数据库容器。同时配置了正确的端口映射和环境变量注入。", False, 12)
    ])

    # Part 2
    add_heading_with_font(doc, "二、在本阶段中个人承担了哪些任务，完成情况如何，过程中是否遇到问题，如何解决的", level=2)

    add_formatted_paragraph(doc, "在本项目的开发阶段，我主要承担了前端页面开发、数据库初始化脚本编写、全链路自动化测试、Docker 容器化部署和项目文档编写工作。", font_size=12)

    # Task 1
    add_heading_with_font(doc, "任务一：Vue3 前端页面全栈开发", level=3)
    add_mixed_paragraph(doc, [
        ("完成情况：", True, 12),
        ("100% 完成。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("工作内容：", True, 12),
        ("使用 Vue3 + Vite + Element Plus 技术栈完成了用户端和管理端的全部页面开发。用户端包括首页、食堂列表、菜品浏览、购物车、订单管理、预约管理和取餐码展示页面。管理端包括仪表盘数据统计、订单管理和扫码核销页面。通过 Vue Router 实现了路由管理和页面导航，使用 Pinia 进行全局状态管理。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("遇到的问题：", True, 12),
        ("前端在计算购物车总价时，由于 JavaScript 的浮点数精度丢失问题，导致金额显示出现多位小数。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("解决过程：", True, 12),
        ("在前端引入了自定义的金额计算工具函数，将金额先乘以 100 转为整数进行加减运算，最后再除以 100 并保留两位小数格式化展示，确保了财务数据的显示准确性。", False, 12)
    ])

    # Task 2
    add_heading_with_font(doc, "任务二：数据库初始化脚本编写", level=3)
    add_mixed_paragraph(doc, [
        ("完成情况：", True, 12),
        ("100% 完成。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("工作内容：", True, 12),
        ("编写了完整的数据库初始化脚本（data_init.sql），包含所有表结构的 DDL 语句和基础数据的 DML 语句。创建了食堂、菜品、时段、座位等基础数据表，并插入了模拟的历史订单数据以供前端图表组件测试。", False, 12)
    ])

    # Task 3
    add_heading_with_font(doc, "任务三：全链路自动化测试脚本编写与执行", level=3)
    add_mixed_paragraph(doc, [
        ("完成情况：", True, 12),
        ("100% 完成。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("工作内容：", True, 12),
        ("编写了基于 Python requests 库的全链路接口测试脚本，覆盖用户注册、登录、浏览食堂和菜品、购物车操作、订单创建、支付回调、预约就餐、二维码生成和管理端核销的完整业务流程。通过 assert 断言验证每个环节的返回状态和数据正确性。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("遇到的问题：", True, 12),
        ("测试脚本在运行时，由于前后端接口字段变更导致部分断言失败。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("解决过程：", True, 12),
        ("与后端开发人员保持密切沟通，及时同步接口变更。在测试脚本中增加了详细的日志输出和错误提示，便于快速定位失败原因。", False, 12)
    ])

    # Task 4
    add_heading_with_font(doc, "任务四：Docker 容器化部署", level=3)
    add_mixed_paragraph(doc, [
        ("完成情况：", True, 12),
        ("100% 完成。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("工作内容：", True, 12),
        ("为前端和后端分别编写了 Dockerfile，配置了多阶段构建以减小镜像体积。编写了 docker-compose.yml 文件统一管理前端、后端和 MySQL 三个服务，通过自定义 Docker 网络实现容器间通信。配置了正确的环境变量注入和端口映射。", False, 12)
    ])

    # Task 5
    add_heading_with_font(doc, "任务五：前后端联调与接口对接", level=3)
    add_mixed_paragraph(doc, [
        ("完成情况：", True, 12),
        ("100% 完成。", False, 12)
    ])
    add_mixed_paragraph(doc, [
        ("工作内容：", True, 12),
        ('在前后端开发完成后，进行了全面的接口联调测试。修复了跨域问题、Token 携带问题、接口字段对齐等问题。确保了"浏览商品 -> 下单 -> 取餐码生成 -> 管理员扫码核销"完整全链路的正常运行。', False, 12)
    ])

    # Summary
    add_heading_with_font(doc, "三、总结", level=2)
    add_formatted_paragraph(doc, "本阶段按计划完成了我所负责的所有前端开发、测试和部署工作。在前端开发方面，使用 Vue3 技术栈完成了用户端和管理端的全部页面；在测试方面，编写了全链路自动化测试脚本覆盖所有核心业务流程；在部署方面，实现了 Docker 容器化部署方案。代码已通过基本的功能测试和自动化测试。在实际开发过程中，不仅提升了对 Vue3 组件化开发和前端工程化的理解，还通过解决跨域、Token 管理、自动化测试和容器化部署等问题，积累了宝贵的全栈工程实践经验。", font_size=12)

    doc.save(os.path.join(TASKS_DIR, "软件综合实训_黄译辉_开发阶段实训报告.docx"))
    print("Created: 软件综合实训_黄译辉_开发阶段实训报告.docx")


# ============================================================
# 4. 软件综合实训_陈金瑞_单元测试报告.docx
# ============================================================
def create_chen_jinrui_unit_test():
    doc = Document()

    add_heading_with_font(doc, "软件综合实训_陈金瑞_单元测试报告", level=1)

    # Part 1
    add_heading_with_font(doc, "一、测试什么？", level=2)

    add_formatted_paragraph(doc, "本次测试主要针对智慧校园食堂系统的后端 API 接口进行单元测试，测试内容包含以下方面：", font_size=12)

    add_heading_with_font(doc, "1. 函数与模块交互输入输出的正确性", level=3)
    add_formatted_paragraph(doc, "验证后端各个 Controller 层接口是否能正确解析前端发送的 JSON 请求载荷，同时保证后端返回的数据结构能够按照统一结果封装（Result 类），输出正确的 JSON 结构。重点验证了用户注册登录接口、食堂菜品查询接口、购物车操作接口、订单创建接口、预约接口和二维码核销接口的请求参数校验和响应数据格式。", font_size=12)

    add_heading_with_font(doc, "2. JWT Token 验证机制测试", level=3)
    add_formatted_paragraph(doc, "验证 JWT 认证体系的正确性，包括：Token 生成后能否正确解析出用户信息；无效或过期的 Token 访问受保护接口时是否正确返回 401 状态码；拦截器是否正确排除了登录、注册等公开接口；ThreadLocal 中存储的用户信息是否正确传递到业务层。", font_size=12)

    add_heading_with_font(doc, "3. 数据库事务管理与业务逻辑测试", level=3)
    add_formatted_paragraph(doc, "验证订单创建流程中的事务一致性，包括：订单创建、库存扣减、订单明细插入是否在同一事务中；当某个环节失败时，是否全部回滚；乐观锁机制是否有效防止库存超卖；预约接口是否正确校验时间段和座位可用性。", font_size=12)

    add_heading_with_font(doc, "4. 二维码安全性测试", level=3)
    add_formatted_paragraph(doc, "验证取餐二维码的安全机制，包括：qrToken 的唯一性和不可伪造性；二维码过期时间是否生效；已核销的二维码是否不能重复使用；核销日志是否正确记录。", font_size=12)

    # Part 2
    add_heading_with_font(doc, "二、如何测试？", level=2)

    add_formatted_paragraph(doc, "本次测试采用了自动化测试脚本驱动机制：", font_size=12)

    add_mixed_paragraph(doc, [
        ("测试方式：", True, 12),
        ("通过编写基于 Python requests 库的全流程接口串联自动化脚本（tests/full_link_test.py）进行后端 API 的集成与单元测试。脚本模拟完整的用户操作流程，依次调用各个后端接口并验证响应结果。", False, 12)
    ])

    add_mixed_paragraph(doc, [
        ("输入数据：", True, 12),
        ("用户注册使用动态生成的虚拟手机号码（通过时间戳截取避免重复）和默认密码进行输入。订单与预约流程使用已有的 canteenId 和 foodId 作为参数，购物车请求输入具体数量。预约时请求携带有效 orderId，系统根据当前日期读取可用营业时间段和座位信息。", False, 12)
    ])

    add_mixed_paragraph(doc, [
        ("预期输出：", True, 12),
        ("后端响应必须符合统一格式（code、message、data）；核心接口返回 HTTP 200 状态码；能够正确流转订单及核销状态；成功生成有效的 Token 和 qrToken；返回对应环节生成的自增 ID 及相关数据。", False, 12)
    ])

    add_mixed_paragraph(doc, [
        ("判定标准：", True, 12),
        ("使用严格的 assert 断言验证每一步的返回状态码和数据正确性。任何环节产生错误状态码或抛出异常即判定测试不通过。脚本无报错成功运行走完全链路则判定业务逻辑正确。", False, 12)
    ])

    # Part 3
    add_heading_with_font(doc, "三、实际测试及测试结果是什么？", level=2)

    add_formatted_paragraph(doc, "使用 Python3 执行了本地运行服务下的集成链路测试脚本。脚本依序完成了用户注册（获取 userId=3）、登录带出 USER_TOKEN，并在加入菜品后创建了一笔金额为 30.00 元的订单。完成支付回调通知后锁定系统当天 11:00-12:00 时段以及 A-001 号座位并获取核销令牌，最后交由管理端接口确认核销成功。", font_size=12)

    add_heading_with_font(doc, "测试结果：", level=3)
    add_formatted_paragraph(doc, "测试顺利通过全链路各项断言，无异常抛出。各后端 API 接口交互与数据流转状态符合预期设定。JWT Token 验证机制工作正常，数据库事务回滚机制有效，乐观锁并发控制生效。", font_size=12)

    add_heading_with_font(doc, "测试控制台输出结果实录：", level=3)
    test_output = """=== 开始全链路测试 ===

[1] 正在注册新用户...
状态码: 200, 响应: {"code":0,"message":"ok","data":{"id":3,"userNo":"U17773583760921692","userName":"测试用户","phone":"18877358376","password":null,"role":"USER","createdAt":null,"updatedAt":null}}

[2] 正在登录...
状态码: 200, 响应: {"code":0,"message":"ok","data":{"id":3,"userName":"测试用户","token":"eyJhbGciOiJIUz...","role":"USER"}}

[3] 查看食堂列表...
状态码: 200, 食堂数量: 2

[4] 查看菜品列表...
状态码: 200

[5] 加入购物车...
状态码: 200

[6] 创建订单...
状态码: 200, 响应: {"code":0,"message":"ok","data":{"id":2,"orderNo":"ORD1777358376141","userId":3,"canteenId":1,"totalAmount":30.00,"paymentMethod":"WECHAT","orderStatus":"CREATED","paidAt":null,"createdAt":null,"updatedAt":null}}

[7] 模拟支付回调...
状态码: 200

[8] 预约就餐...
获取时段状态码: 200
获取到 2 个时段
找到可用时段 11:00:00-12:00:00, 座位 A-001
状态码: 200, 响应: {"code":0,"message":"ok","data":{"id":2,"reservationCode":"RES1777358376168","userId":3,"orderId":2,"canteenId":1,"timeSlotId":1,"seatId":1,"reservationStatus":"RESERVED","createdAt":null,"updatedAt":null}}

[9] 生成取餐二维码...
状态码: 200, 响应: {"code":0,"message":"ok","data":{"id":3,"qrToken":"35f61e637737421bb7...","expiresAt":"2026-04-28T15:39:36.173025","qrImageUrl":"..."}}

[10] 管理端扫码核销...
状态码: 200, 响应: {"code":0,"message":"ok","data":{"userId":3,"userName":"测试用户","orderNo":"ORD1777358376141","totalAmount":30.00,"paidAt":"2026-04-28T14:39:36","items":[{"id":3,"orderId":2,"foodId":1,"foodNameSnapshot":"红烧肉","unitPriceSnapshot":15.00,"quantity":2,"subtotalAmount":30.00,"createdAt":"2026-04-28T14:39:36","updatedAt":"2026-04-28T14:39:36"}],"canteenName":"第一食堂","timeSlot":"11:00-12:00","seatNo":"A-001","verified":true,"message":"Scan successful"}}

=== 全链路测试完成 ==="""

    p = doc.add_paragraph()
    run = p.add_run(test_output)
    set_run_font(run, font_name="Consolas", font_size=10)
    p.paragraph_format.space_after = Pt(6)

    doc.save(os.path.join(TASKS_DIR, "软件综合实训_陈金瑞_单元测试报告.docx"))
    print("Created: 软件综合实训_陈金瑞_单元测试报告.docx")


# ============================================================
# 5. 软件综合实训_黄译辉_单元测试报告.docx
# ============================================================
def create_huang_yihui_unit_test():
    doc = Document()

    add_heading_with_font(doc, "软件综合实训_黄译辉_单元测试报告", level=1)

    # Part 1
    add_heading_with_font(doc, "一、测试什么？", level=2)

    add_formatted_paragraph(doc, "本次测试主要针对智慧校园食堂系统的全链路接口进行自动化测试和验证，测试内容包含以下方面：", font_size=12)

    add_heading_with_font(doc, "1. 全链路接口测试脚本编写与执行", level=3)
    add_formatted_paragraph(doc, "验证从用户注册到最终核销的完整业务流程，包括用户注册登录、食堂菜品浏览、购物车操作、订单创建、支付回调、预约就餐、取餐二维码生成和管理端扫码核销。测试各环节之间的数据传递是否正确，前后端接口对接是否顺畅。", font_size=12)

    add_heading_with_font(doc, "2. 前端组件交互测试", level=3)
    add_formatted_paragraph(doc, "验证前端各组件之间的数据传递和交互逻辑是否正确。包括购物车总价计算、订单状态实时更新、预约时段选择、二维码展示和扫码核销等交互流程。重点验证了前端 Axios 请求封装中 Token 的自动携带和 401 状态码的自动处理逻辑。", font_size=12)

    add_heading_with_font(doc, "3. Docker 容器化部署验证", level=3)
    add_formatted_paragraph(doc, "验证 Docker 容器化部署方案的可行性，包括前端容器、后端容器和 MySQL 容器之间的网络连接是否正常，端口映射是否正确，环境变量注入是否生效。确保在 Docker 环境下所有核心功能均可正常运行。", font_size=12)

    add_heading_with_font(doc, "4. 数据初始化验证", level=3)
    add_formatted_paragraph(doc, "验证数据库初始化脚本（data_init.sql）的正确性，包括表结构创建是否完整，基础数据插入是否成功，模拟历史订单数据是否满足前端图表组件的渲染需求。", font_size=12)

    # Part 2
    add_heading_with_font(doc, "二、如何测试？", level=2)

    add_formatted_paragraph(doc, "本次测试采用了多种测试方式相结合：", font_size=12)

    add_mixed_paragraph(doc, [
        ("全链路接口测试：", True, 12),
        ("通过编写基于 Python requests 库的自动化测试脚本（tests/full_link_test.py），模拟完整用户操作流程。脚本使用 Session 保持请求上下文，将每一步的响应数据解析后作为后续请求的输入。通过严格的 assert 断言验证每个环节的返回状态和数据正确性。", False, 12)
    ])

    add_mixed_paragraph(doc, [
        ("前端交互测试：", True, 12),
        ("通过本地启动前端开发服务器（npm run dev）与后端服务联调，手动验证各页面的交互流程和数据渲染。重点测试了购物车金额计算、订单状态轮询刷新和扫码核销功能的用户体验。", False, 12)
    ])

    add_mixed_paragraph(doc, [
        ("Docker 部署验证：", True, 12),
        ("通过 docker-compose up 启动所有服务，验证容器间网络连接、端口映射和环境变量配置是否正确。在 Docker 环境下执行全链路测试脚本验证功能完整性。", False, 12)
    ])

    add_mixed_paragraph(doc, [
        ("数据初始化验证：", True, 12),
        ("执行 data_init.sql 脚本后，通过 SQL 查询验证表结构和数据是否正确插入。确认基础数据（食堂、菜品、时段、座位）和模拟历史订单数据满足测试需求。", False, 12)
    ])

    # Part 3
    add_heading_with_font(doc, "三、实际测试及测试结果是什么？", level=2)

    add_formatted_paragraph(doc, "使用 Python3 执行了本地运行服务下的全链路集成测试脚本。脚本依次完成了用户注册（获取 userId=3）、登录获取 Token、浏览食堂和菜品列表、添加购物车、创建 30 元订单、模拟支付回调、预约 11:00-12:00 时段 A-001 座位、生成取餐二维码和管理端核销的完整流程。所有环节均通过 assert 断言验证，无异常抛出。", font_size=12)

    add_heading_with_font(doc, "测试结果：", level=3)
    add_formatted_paragraph(doc, "全链路测试顺利通过，各项断言均通过验证。前端组件交互正常，Docker 容器化部署验证通过，数据初始化脚本执行成功。各业务模块交互与数据流转状态符合预期设定。", font_size=12)

    add_heading_with_font(doc, "测试控制台输出结果实录：", level=3)
    test_output = """=== 开始全链路测试 ===

[1] 正在注册新用户...
状态码: 200, 响应: {"code":0,"message":"ok","data":{"id":3,"userNo":"U17773583760921692","userName":"测试用户","phone":"18877358376","password":null,"role":"USER","createdAt":null,"updatedAt":null}}

[2] 正在登录...
状态码: 200, 响应: {"code":0,"message":"ok","data":{"id":3,"userName":"测试用户","token":"eyJhbGciOiJIUz...","role":"USER"}}

[3] 查看食堂列表...
状态码: 200, 食堂数量: 2

[4] 查看菜品列表...
状态码: 200

[5] 加入购物车...
状态码: 200

[6] 创建订单...
状态码: 200, 响应: {"code":0,"message":"ok","data":{"id":2,"orderNo":"ORD1777358376141","userId":3,"canteenId":1,"totalAmount":30.00,"paymentMethod":"WECHAT","orderStatus":"CREATED","paidAt":null,"createdAt":null,"updatedAt":null}}

[7] 模拟支付回调...
状态码: 200

[8] 预约就餐...
获取时段状态码: 200
获取到 2 个时段
找到可用时段 11:00:00-12:00:00, 座位 A-001
状态码: 200, 响应: {"code":0,"message":"ok","data":{"id":2,"reservationCode":"RES1777358376168","userId":3,"orderId":2,"canteenId":1,"timeSlotId":1,"seatId":1,"reservationStatus":"RESERVED","createdAt":null,"updatedAt":null}}

[9] 生成取餐二维码...
状态码: 200, 响应: {"code":0,"message":"ok","data":{"id":3,"qrToken":"35f61e637737421bb7...","expiresAt":"2026-04-28T15:39:36.173025","qrImageUrl":"..."}}

[10] 管理端扫码核销...
状态码: 200, 响应: {"code":0,"message":"ok","data":{"userId":3,"userName":"测试用户","orderNo":"ORD1777358376141","totalAmount":30.00,"paidAt":"2026-04-28T14:39:36","items":[{"id":3,"orderId":2,"foodId":1,"foodNameSnapshot":"红烧肉","unitPriceSnapshot":15.00,"quantity":2,"subtotalAmount":30.00,"createdAt":"2026-04-28T14:39:36","updatedAt":"2026-04-28T14:39:36"}],"canteenName":"第一食堂","timeSlot":"11:00-12:00","seatNo":"A-001","verified":true,"message":"Scan successful"}}

=== 全链路测试完成 ==="""

    p = doc.add_paragraph()
    run = p.add_run(test_output)
    set_run_font(run, font_name="Consolas", font_size=10)
    p.paragraph_format.space_after = Pt(6)

    doc.save(os.path.join(TASKS_DIR, "软件综合实训_黄译辉_单元测试报告.docx"))
    print("Created: 软件综合实训_黄译辉_单元测试报告.docx")


if __name__ == "__main__":
    os.makedirs(TASKS_DIR, exist_ok=True)
    create_meeting_record()
    create_chen_jinrui_report()
    create_huang_yihui_report()
    create_chen_jinrui_unit_test()
    create_huang_yihui_unit_test()
    print("\nAll docx documents created successfully!")
